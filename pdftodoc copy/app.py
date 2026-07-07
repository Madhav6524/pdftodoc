import os, io, re, tempfile, traceback, uuid, threading

# Load .env early so OPENAI_API_KEY is available everywhere
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

import numpy as np
from flask import Flask, request, send_file, jsonify, render_template
from flask_cors import CORS
import fitz   # PyMuPDF

try:
    from pdf2docx import Converter as PDFConverter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as _Paragraph
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

app = Flask(__name__, static_folder="static", static_url_path="/static")
CORS(app)

# ─────────────────────────────────────────────────────────────────────────────
# Font helpers (text-layer replacement)
# ─────────────────────────────────────────────────────────────────────────────

def _clean(name):
    return re.sub(r'^[A-Z]{6}\+', '', name or '')

def _unpack_color(packed):
    return (((packed>>16)&0xFF)/255.0, ((packed>>8)&0xFF)/255.0, (packed&0xFF)/255.0)

def _pick_builtin(name, flags):
    bold   = bool(flags & 0b10000) or bool(re.search(r'bold',          name, re.I))
    italic = bool(flags & 0b00010) or bool(re.search(r'italic|oblique', name, re.I))
    serif  = bool(flags & 0b00100)
    nl = name.lower()
    if 'courier' in nl or 'mono' in nl:
        return 'cobo' if bold else ('coit' if italic else 'cour')
    if serif or 'times' in nl or 'roman' in nl or 'serif' in nl:
        return 'tibo' if (bold or italic) else 'tiro'
    if bold and italic: return 'hebo'
    if bold:            return 'hebo'
    if italic:          return 'heit'
    return 'helv'

_MAC_FONT_PATHS = [
    ("arial bold italic",  ["/Library/Fonts/Arial Bold Italic.ttf"]),
    ("arial bold",         ["/Library/Fonts/Arial Bold.ttf"]),
    ("arial italic",       ["/Library/Fonts/Arial Italic.ttf"]),
    ("arial",              ["/Library/Fonts/Arial.ttf",
                             "/System/Library/Fonts/Supplemental/Arial.ttf"]),
    ("timesnewroman bold", ["/Library/Fonts/Times New Roman Bold.ttf"]),
    ("timesnewroman",      ["/Library/Fonts/Times New Roman.ttf",
                             "/System/Library/Fonts/Supplemental/Times New Roman.ttf"]),
    ("calibri bold",       ["/Library/Fonts/Microsoft/Calibri Bold.ttf"]),
    ("calibri",            ["/Library/Fonts/Microsoft/Calibri.ttf",
                             "/Library/Fonts/Calibri.ttf"]),
    ("helveticaneue",      ["/System/Library/Fonts/HelveticaNeue.ttc"]),
    ("helvetica",          ["/System/Library/Fonts/Helvetica.ttc"]),
    ("georgia",            ["/Library/Fonts/Georgia.ttf"]),
    ("verdana",            ["/Library/Fonts/Verdana.ttf"]),
    ("trebuchet",          ["/Library/Fonts/Trebuchet MS.ttf"]),
    ("tahoma",             ["/Library/Fonts/Tahoma.ttf"]),
    ("garamond",           ["/Library/Fonts/Garamond.ttf"]),
    ("courier",            ["/Library/Fonts/Courier New.ttf"]),
]

def _load_font(doc, page, font_name, flags):
    clean  = _clean(font_name)
    target = clean.lower()
    for fi in page.get_fonts(full=True):
        xref = fi[0]
        if not xref: continue
        bf = _clean(fi[3]).lower()
        fn = _clean(fi[4]).lower()
        if target == bf or target == fn or target in bf or target in fn:
            try:
                fd = doc.extract_font(xref)
                if fd and len(fd) >= 4 and fd[3] and len(fd[3]) > 100:
                    font = fitz.Font(fontbuffer=fd[3])
                    return font, True
            except Exception as e:
                print(f"[FONT] Embedded extract failed: {e}")

    name_key = re.sub(r'[^a-z]', '', clean.lower())
    for key, paths in _MAC_FONT_PATHS:
        key_norm = key.replace(' ', '')
        if key_norm in name_key or name_key in key_norm:
            for path in paths:
                if os.path.exists(path):
                    try:
                        return fitz.Font(fontfile=path), True
                    except: pass

    return None, False

def _font_covers_text(font_obj, text):
    """
    Return True only if every non-space character in `text` has a real glyph
    in this font object. Embedded PDF fonts are almost always *subsetted* —
    they only contain glyphs for characters that appeared in the original
    document. If we blindly trust an embedded bold/serif font object just
    because it loaded successfully, PyMuPDF's TextWriter will silently
    substitute missing glyphs with a fallback face, which is why replacement
    text (e.g. "Patel") can render in the wrong weight/style even though the
    rest of the line looks correct. This check forces a fallback to a
    style-matched builtin font whenever the embedded subset can't render the
    replacement word.
    """
    if font_obj is None:
        return False
    try:
        for ch in text:
            if ch.isspace():
                continue
            if not font_obj.has_glyph(ord(ch)):
                return False
        return True
    except Exception:
        return False

# ─────────────────────────────────────────────────────────────────────────────
# Text-layer replacement
# ─────────────────────────────────────────────────────────────────────────────

def replace_text_in_pdf(doc, find_word, replace_word):
    total = 0
    for page in doc:
        raw = page.get_text("rawdict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        hits = []

        for blk in raw.get("blocks", []):
            if blk.get("type") != 0:
                continue
            for line in blk.get("lines", []):
                line_spans = line.get("spans", [])
                if not line_spans:
                    continue

                # ── Determine the dominant font/style for this whole line ──
                # PDFs (especially form-style templates) sometimes store one
                # word in a different embedded font than the rest of the
                # line (e.g. a filled-in name vs. the surrounding label
                # text). If we blindly reuse the matched span's own font,
                # the replacement can visually clash with its neighbors.
                # Instead, vote by character count across all spans on the
                # line and use whichever font/style wins, so the replaced
                # word matches what the eye actually sees as "the line's
                # font" rather than a possibly-mismatched per-span quirk.
                font_votes = {}
                for sp in line_spans:
                    chars = sp.get("chars", [])
                    if not chars:
                        continue
                    key = (sp.get("font", "Helvetica"), int(sp.get("flags", 0)))
                    font_votes[key] = font_votes.get(key, 0) + len(chars)

                if font_votes:
                    dominant_font, dominant_flags = max(
                        font_votes.items(), key=lambda kv: kv[1]
                    )[0]
                else:
                    dominant_font, dominant_flags = ("Helvetica", 0)

                line_has_mixed_fonts = len(font_votes) > 1

                for span in line_spans:
                    chars = span.get("chars", [])
                    if not chars:
                        continue
                    span_text = "".join(c.get("c", "") for c in chars)
                    fn_orig    = span.get("font", "Helvetica")
                    fs         = float(span.get("size", 11))
                    color      = int(span.get("color", 0))
                    flags_orig = int(span.get("flags", 0))

                    idx = 0
                    while True:
                        pos = span_text.lower().find(find_word.lower(), idx)
                        if pos == -1:
                            break
                        mc = chars[pos:pos + len(find_word)]
                        if not mc:
                            break
                        x0 = mc[0]["bbox"][0]; y0 = min(c["bbox"][1] for c in mc)
                        x1 = mc[-1]["bbox"][2]; y1 = max(c["bbox"][3] for c in mc)

                        # Use this span's own font only if the whole line is
                        # already uniform. If the line mixes styles, trust
                        # the dominant one instead of the matched span's.
                        if line_has_mixed_fonts:
                            use_font, use_flags = dominant_font, dominant_flags
                        else:
                            use_font, use_flags = fn_orig, flags_orig

                        hits.append((
                            fitz.Rect(x0, y0, x1, y1),
                            fitz.Point(mc[0]["origin"]),
                            use_font, fs, color, use_flags
                        ))
                        idx = pos + 1

        if not hits:
            continue

        fill_colors = []
        for rect, *_ in hits:
            fill = (1.0, 1.0, 1.0)
            try:
                pm = page.get_pixmap(clip=rect.irect, matrix=fitz.Matrix(1, 1), colorspace=fitz.csRGB)
                if pm.samples:
                    fill = (pm.samples[0] / 255.0, pm.samples[1] / 255.0, pm.samples[2] / 255.0)
            except:
                pass
            fill_colors.append(fill)

        for (rect, *_), fill in zip(hits, fill_colors):
            page.add_redact_annot(rect, fill=fill)
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        for (rect, baseline, fn, fs, color_int, flags), fill in zip(hits, fill_colors):
            text_color = _unpack_color(color_int)
            font_obj, is_real = _load_font(doc, page, fn, flags)
            if is_real and font_obj and _font_covers_text(font_obj, replace_word):
                try:
                    tw = fitz.TextWriter(page.rect)
                    tw.append(baseline, replace_word, font=font_obj, fontsize=fs)
                    tw.write_text(page, color=text_color)
                    total += 1
                    continue
                except Exception as e:
                    print(f"[WARN] TextWriter failed: {e}")
            page.insert_text(baseline, replace_word,
                             fontname=_pick_builtin(fn, flags),
                             fontsize=fs, color=text_color)
            total += 1
    return total

# ─────────────────────────────────────────────────────────────────────────────
# In-memory download store
# ─────────────────────────────────────────────────────────────────────────────

_store      = {}
_store_lock = threading.Lock()

def _store_put(data: bytes, mime: str, filename: str) -> str:
    token = str(uuid.uuid4())
    with _store_lock:
        _store[token] = {"data": data, "mime": mime, "filename": filename}
    return token

def _store_pop(token: str):
    with _store_lock:
        return _store.pop(token, None)

# ─────────────────────────────────────────────────────────────────────────────
# Routes
# ─────────────────────────────────────────────────────────────────────────────

@app.route("/replace", methods=["POST"])
def replace():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file_obj = request.files["file"]
    if not file_obj.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Must be a PDF"}), 400

    out_fmt     = request.form.get("output_format", "pdf").strip().lower()
    output_name = request.form.get("output_name", "").strip()

    # Support multiple pairs via JSON field "pairs"
    import json as _json
    pairs_raw = request.form.get("pairs", "")
    try:
        pairs = _json.loads(pairs_raw) if pairs_raw else []
    except Exception:
        pairs = []

    # Fallback: single pair from legacy fields
    if not pairs:
        fw = request.form.get("find_word",    "").strip()
        rw = request.form.get("replace_word", "").strip()
        if fw:
            pairs = [{"find": fw, "replace": rw}]

    pairs = [p for p in pairs if p.get("find","").strip()]
    if not pairs:
        return jsonify({"error": "At least one find/replace pair is required"}), 400
    if out_fmt == "docx" and not PDF2DOCX_AVAILABLE:
        return jsonify({"error": "pdf2docx not installed"}), 500

    pdf_bytes  = file_obj.read()
    doc        = fitz.open(stream=pdf_bytes, filetype="pdf")

    text_total   = 0
    pair_results = []

    for pair in pairs:
        fw = pair.get("find", "").strip()
        rw = pair.get("replace", "").strip()
        if not fw: continue
        tc = replace_text_in_pdf(doc, fw, rw)
        text_total += tc
        pair_results.append({"find": fw, "replace": rw, "total": tc})

    out_buf = io.BytesIO()
    doc.save(out_buf, garbage=4, deflate=True)
    doc.close()
    out_buf.seek(0)
    pdf_out = out_buf.getvalue()

    # Determine output filename
    base = os.path.splitext(file_obj.filename)[0]
    if output_name:
        out_stem = re.sub(r'[^\w\s\-]', '', output_name).strip() or base
    else:
        out_stem = f"{base}_modified"

    total = text_total

    if out_fmt == "docx":
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
            tf.write(pdf_out); tmp_pdf = tf.name
        tmp_docx = tmp_pdf.replace(".pdf", ".docx")
        try:
            cv = PDFConverter(tmp_pdf)
            cv.convert(tmp_docx, start=0, end=None)
            cv.close()
            with open(tmp_docx, "rb") as f:
                docx_bytes = f.read()
        finally:
            for p in (tmp_pdf, tmp_docx):
                try: os.unlink(p)
                except: pass
        fname_docx = f"{out_stem}.docx"
        token = _store_put(docx_bytes,
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           fname_docx)
        return jsonify({"token": token, "filename": fname_docx,
                        "text_count": text_total,
                        "total": total,
                        "pairs": pair_results})

    fname_pdf = f"{out_stem}.pdf"
    token = _store_put(pdf_out, "application/pdf", fname_pdf)
    return jsonify({"token": token, "filename": fname_pdf,
                    "text_count": text_total,
                    "total": total,
                    "pairs": pair_results})

@app.route("/merge", methods=["POST"])
def merge_pdfs():
    """Merge multiple uploaded PDFs into one."""
    files = request.files.getlist("files")
    if len(files) < 2:
        return jsonify({"error": "Please upload at least 2 PDF files"}), 400

    output_name = request.form.get("output_name", "merged").strip()
    # Sanitise filename
    output_name = re.sub(r'[^\w\s\-]', '', output_name).strip() or "merged"
    if not output_name.lower().endswith(".pdf"):
        output_name += ".pdf"

    merged = fitz.open()
    errors = []
    for f in files:
        if not f.filename.lower().endswith(".pdf"):
            errors.append(f"{f.filename}: not a PDF")
            continue
        try:
            src = fitz.open(stream=f.read(), filetype="pdf")
            merged.insert_pdf(src)
            src.close()
        except Exception as e:
            errors.append(f"{f.filename}: {e}")

    if merged.page_count == 0:
        return jsonify({"error": "No valid PDFs could be merged"}), 400

    # Capture page count BEFORE closing
    total_pages = merged.page_count
    buf = io.BytesIO()
    merged.save(buf, garbage=4, deflate=True)
    merged.close()
    buf.seek(0)

    token = _store_put(buf.getvalue(), "application/pdf", output_name)
    return jsonify({"token": token, "pages": total_pages,
                    "errors": errors,
                    "filename": output_name})

@app.route("/convert-to-docx", methods=["POST"])
def convert_to_docx():
    """Convert PDF to DOCX. Default is editable pdf2docx; visual mode is opt-in."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f = request.files["file"]
    if not f.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Must be a PDF file"}), 400

    pdf_bytes = f.read()
    layout_mode = request.form.get("layout_mode", "editable").strip().lower()

    if layout_mode == "visual":
        try:
            docx_bytes = _convert_pdf_to_docx_visual(pdf_bytes)
        except Exception as visual_err:
            print(f"[DOCX] Visual conversion failed, falling back: {visual_err}")
            traceback.print_exc()
            if not PDF2DOCX_AVAILABLE:
                return jsonify({"error": f"Visual conversion failed and pdf2docx is not installed: {visual_err}"}), 500
            try:
                docx_bytes = _convert_pdf_to_docx_pdf2docx(pdf_bytes)
            except Exception as e:
                traceback.print_exc()
                return jsonify({"error": str(e)}), 500
    else:
        if not PDF2DOCX_AVAILABLE:
            return jsonify({"error": "pdf2docx not installed"}), 500
        try:
            docx_bytes = _convert_pdf_to_docx_pdf2docx(pdf_bytes)
        except Exception as e:
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    base  = os.path.splitext(f.filename)[0]
    token = _store_put(
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{base}.docx"
    )
    return jsonify({"token": token, "filename": f"{base}.docx"})

def _convert_pdf_to_docx_pdf2docx(pdf_bytes):
    """Primary editable conversion path: header-wrap repair + real Word
    header/footer handling (including live PAGE/NUMPAGES fields)."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
        tf.write(pdf_bytes); tmp_pdf = tf.name
    tmp_docx = tmp_pdf.replace(".pdf", ".docx")

    try:
        cv = PDFConverter(tmp_pdf)
        cv.convert(tmp_docx, start=0, end=None)
        cv.close()

        if PYTHON_DOCX_AVAILABLE:
            _repair_docx_header_wrapping(tmp_docx, pdf_bytes)

            try:
                hdr_info, ftr_info = _extract_header_footer_info(pdf_bytes)
                if hdr_info or ftr_info:
                    _apply_headers_footers(tmp_docx, hdr_info, ftr_info)
            except Exception as e:
                print(f"[DOCX] Header/footer handling failed: {e}")
                traceback.print_exc()

        with open(tmp_docx, "rb") as fh:
            docx_bytes = fh.read()
    finally:
        for p in (tmp_pdf, tmp_docx):
            try: os.unlink(p)
            except: pass
    return docx_bytes

def _iter_body_paragraphs(document):
    """
    Yield every paragraph in the document body, IN DOCUMENT ORDER, including
    paragraphs nested inside tables (at any depth).

    python-docx's `document.paragraphs` only returns top-level body
    paragraphs and silently skips anything inside a table. pdf2docx commonly
    rebuilds multi-column PDF headers/footers (e.g. "Subject Code | Subject
    Name" on one row, "Enrollment No | Name" on the next) as a Word table to
    preserve the column positions. If we only ever look at
    `document.paragraphs`, that fragmented header text sitting inside the
    table cells is invisible to both the header-wrap repair and the
    duplicate-removal step — which is why it kept showing up as leftover
    body text even after the real header was inserted correctly. Walking
    every <w:p> element in the body (tables included) fixes that.
    """
    body = document.element.body
    for p in body.iter(qn('w:p')):
        yield _Paragraph(p, document)

def _remove_empty_tables(document):
    """Delete any table (including nested tables) whose cells are all
    blank after duplicate header/footer text has been stripped out of it."""
    def cells_all_empty(table):
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    return False
        return True

    def process(tables):
        for table in list(tables):
            for row in table.rows:
                for cell in row.cells:
                    process(cell.tables)
            if cells_all_empty(table):
                tbl_el = table._tbl
                parent = tbl_el.getparent()
                if parent is not None:
                    parent.remove(tbl_el)

    process(document.tables)

def _repair_docx_header_wrapping(docx_path, pdf_bytes):
    """
    pdf2docx sometimes fragments a PDF header line into many Word paragraphs
    (including paragraphs nested inside a table, when it used a table to
    preserve a multi-column header layout). Rebuild only lines that are
    known to be in the PDF header zone, leaving other body content alone.
    """
    header_lines = _extract_pdf_header_lines(pdf_bytes)
    if not header_lines:
        return

    docx = DocxDocument(docx_path)
    all_paragraphs = list(_iter_body_paragraphs(docx))
    merges = 0
    for expected in header_lines:
        merges += _merge_fragmented_paragraph_text(all_paragraphs, expected)

    if merges:
        docx.save(docx_path)
        print(f"[DOCX] Repaired {merges} fragmented header line(s)")

def _extract_pdf_header_lines(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    lines = []
    try:
        for page in doc:
            header_limit = page.rect.height * 0.18
            for line in _extract_pdf_lines(page, y_limit=header_limit):
                text = line["text"].strip()
                if not text or len(text) < 2:
                    continue
                norm = _norm_join_text(text)
                if len(norm) < 2:
                    continue
                if all(_norm_join_text(existing) != norm for existing in lines):
                    lines.append(text)
            if len(lines) >= 8:
                break
    finally:
        doc.close()
    return lines[:8]

def _extract_pdf_lines(page, y_limit=None):
    pdf_lines = []
    for block in page.get_text("dict").get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            bbox = line.get("bbox", (0, 0, 0, 0))
            if y_limit is not None and bbox[1] > y_limit:
                continue
            spans = sorted(line.get("spans", []), key=lambda s: s.get("bbox", (0, 0, 0, 0))[0])
            text = "".join(span.get("text", "") for span in spans)
            if text.strip():
                pdf_lines.append({"text": text, "bbox": bbox})
    pdf_lines.sort(key=lambda x: (x["bbox"][1], x["bbox"][0]))
    return pdf_lines

def _merge_fragmented_paragraph_text(paragraphs, expected_text, max_parts=10):
    expected_norm = _norm_join_text(expected_text)
    if len(expected_norm) < 2:
        return 0

    merges = 0
    i = 0
    while i < len(paragraphs):
        if not paragraphs[i].text.strip():
            i += 1
            continue

        acc_norm = ""
        parts = []
        j = i
        while j < len(paragraphs) and len(parts) < max_parts:
            text = paragraphs[j].text.strip()
            if text:
                parts.append(j)
                acc_norm += _norm_join_text(text)
                if acc_norm == expected_norm:
                    if len(parts) > 1:
                        _replace_paragraph_text(paragraphs[parts[0]], expected_text)
                        for idx in reversed(parts[1:]):
                            _delete_paragraph(paragraphs[idx])
                        merges += 1
                        i = parts[0] + 1
                    break
                if not expected_norm.startswith(acc_norm):
                    break
            j += 1
        i += 1
    return merges

def _norm_join_text(text):
    return re.sub(r"[^a-z0-9]+", "", (text or "").lower())

def _replace_paragraph_text(para, text):
    first_run = para.runs[0] if para.runs else None
    para.clear()
    run = para.add_run(text)
    if first_run is not None:
        run.bold = first_run.bold
        run.italic = first_run.italic
        run.underline = first_run.underline
        run.font.size = first_run.font.size
        run.font.name = first_run.font.name
        if first_run.font.color and first_run.font.color.rgb:
            run.font.color.rgb = first_run.font.color.rgb
    fmt = para.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)

def _delete_paragraph(para):
    el = para._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

PT_TO_EMU = 12700
PT_PER_INCH = 72.0
FULL_PAGE_IMAGE_INSET_PT = 1.0

def _convert_pdf_to_docx_visual(pdf_bytes):
    """
    Preserve page count and visual layout by rendering each PDF page as one
    full-page image in the DOCX. This avoids Word text reflow changing headers,
    image positions, spacing, and page breaks.
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")

    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    if pdf.page_count == 0:
        pdf.close()
        raise RuntimeError("PDF has no pages")

    docx = DocxDocument()
    try:
        _set_visual_doc_defaults(docx)
        first_page = pdf[0]
        _set_visual_section(docx.sections[0], first_page)

        for page_index, page in enumerate(pdf):
            if page_index > 0:
                docx.add_page_break()
            para = docx.add_paragraph()
            _prepare_zero_spacing_paragraph(para)

            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), colorspace=fitz.csRGB, alpha=False)
            width_pt = max(1.0, float(page.rect.width) - FULL_PAGE_IMAGE_INSET_PT)
            height_pt = max(1.0, float(page.rect.height) - FULL_PAGE_IMAGE_INSET_PT)
            run = para.add_run()
            run.add_picture(
                io.BytesIO(pix.tobytes("png")),
                width=Inches(width_pt / PT_PER_INCH),
                height=Inches(height_pt / PT_PER_INCH),
            )
    finally:
        pdf.close()

    out = io.BytesIO()
    docx.save(out)
    out.seek(0)
    return out.getvalue()

def _set_visual_doc_defaults(docx):
    style = docx.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = Pt(1)
    style.font.size = Pt(1)

def _set_visual_section(section, page):
    section.page_width = Pt(page.rect.width)
    section.page_height = Pt(page.rect.height)
    section.left_margin = Pt(0)
    section.top_margin = Pt(0)
    section.right_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)

def _prepare_zero_spacing_paragraph(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = para.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(1)

def _convert_pdf_to_docx_hybrid(pdf_bytes):
    """
    Rebuild selectable text + embedded image pages with python-docx while
    deriving all vertical spacing from PDF bounding boxes.
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")

    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    if pdf.page_count == 0:
        pdf.close()
        raise RuntimeError("PDF has no pages")

    docx = DocxDocument()
    prev_block = None
    try:
        for page_index, page in enumerate(pdf):
            page_blocks = _extract_page_layout_blocks(pdf, page)
            if not page_blocks:
                continue
            if page_index == 0:
                _set_docx_page_layout(docx.sections[0], page, page_blocks, image_only=_is_image_only_page(page_blocks))
            if _is_image_only_page(page_blocks):
                _add_page_snapshot(docx, page, prev_block)
                prev_block = {"bbox": (0, 0, page.rect.width, page.rect.height)}
            else:
                prev_block = _write_layout_blocks(docx, page, page_blocks, prev_block)
            if page_index < pdf.page_count - 1:
                docx.add_page_break()
                prev_block = None
    finally:
        pdf.close()

    out = io.BytesIO()
    docx.save(out)
    out.seek(0)
    return out.getvalue()

def _extract_page_layout_blocks(pdf, page):
    text_blocks = _extract_text_line_blocks(page)
    image_blocks = _extract_image_blocks(pdf, page)
    blocks = text_blocks + image_blocks
    blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0], 0 if b["kind"] == "image" else 1))
    return blocks

def _extract_text_line_blocks(page):
    blocks = []
    text_dict = page.get_text("dict")
    for block in text_dict.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = []
            line_text = ""
            for span in line.get("spans", []):
                text = span.get("text", "")
                if not text:
                    continue
                spans.append({
                    "text": text,
                    "font": span.get("font", ""),
                    "size": float(span.get("size", 11)),
                    "bold": bool(int(span.get("flags", 0)) & 0b10000),
                    "italic": bool(int(span.get("flags", 0)) & 0b00010),
                    "color": int(span.get("color", 0)),
                })
                line_text += text
            if not line_text.strip():
                continue
            bbox = tuple(float(v) for v in line.get("bbox", block.get("bbox", (0, 0, 0, 0))))
            blocks.append({
                "kind": "text",
                "bbox": bbox,
                "text": line_text,
                "spans": spans,
            })
    return blocks

def _extract_image_blocks(pdf, page):
    blocks = []
    seen = set()
    for img in page.get_images(full=True):
        xref = img[0]
        try:
            rects = page.get_image_rects(xref)
        except Exception:
            rects = []
        if not rects:
            continue
        try:
            pix = fitz.Pixmap(pdf, xref)
            if pix.n > 4:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            img_bytes = pix.tobytes("png")
        except Exception:
            traceback.print_exc()
            continue
        for rect in rects:
            key = (xref, round(rect.x0, 2), round(rect.y0, 2), round(rect.x1, 2), round(rect.y1, 2))
            if key in seen:
                continue
            seen.add(key)
            blocks.append({
                "kind": "image",
                "bbox": (float(rect.x0), float(rect.y0), float(rect.x1), float(rect.y1)),
                "data": img_bytes,
                "width_pt": float(rect.width),
                "height_pt": float(rect.height),
                "xref": xref,
            })
    return blocks

def _set_docx_page_layout(section, page, blocks, image_only=False):
    section.page_width = Pt(page.rect.width)
    section.page_height = Pt(page.rect.height)

    if image_only:
        section.left_margin = Pt(0)
        section.top_margin = Pt(0)
        section.right_margin = Pt(0)
        section.bottom_margin = Pt(0)
        return

    xs = [b["bbox"][0] for b in blocks] + [b["bbox"][2] for b in blocks]
    ys = [b["bbox"][1] for b in blocks] + [b["bbox"][3] for b in blocks]
    left = max(0.0, min(xs))
    top = max(0.0, min(ys))
    right = max(0.0, page.rect.width - max(xs))
    bottom = max(0.0, page.rect.height - max(ys))

    section.left_margin = Pt(left)
    section.top_margin = Pt(top)
    section.right_margin = Pt(right)
    section.bottom_margin = Pt(bottom)

def _is_image_only_page(blocks):
    return any(b["kind"] == "image" for b in blocks) and not any(b["kind"] == "text" for b in blocks)

def _write_layout_blocks(docx, page, blocks, prev_block):
    text_blocks = [b for b in blocks if b["kind"] == "text"]
    consumed = set()
    for idx, block in enumerate(blocks):
        if idx in consumed:
            continue
        if block["kind"] == "image" and _is_float_image(block, text_blocks):
            side_text = [
                (i, tb) for i, tb in enumerate(blocks)
                if tb["kind"] == "text" and i not in consumed and _y_overlaps(block["bbox"], tb["bbox"])
            ]
            if side_text:
                _add_side_by_side_table(docx, page, block, [tb for _, tb in side_text], prev_block)
                consumed.add(idx)
                for i, _ in side_text:
                    consumed.add(i)
                prev_block = _union_blocks([block] + [tb for _, tb in side_text])
                continue
        if block["kind"] == "text":
            _add_text_block(docx, block, prev_block)
        elif block["kind"] == "image":
            _add_image_block(docx, page, block, prev_block)
        prev_block = block
    return prev_block

def _is_float_image(img_block, text_blocks):
    return any(_y_overlaps(img_block["bbox"], tb["bbox"]) for tb in text_blocks)

def _y_overlaps(a, b):
    return a[1] < b[3] and b[1] < a[3]

def _gap_to_emu(prev_block, curr_block):
    if not prev_block:
        return 0
    gap_pt = float(curr_block["bbox"][1]) - float(prev_block["bbox"][3])
    return int(max(0.0, gap_pt) * PT_TO_EMU)

def _add_text_block(container, block, prev_block=None):
    para = container.add_paragraph()
    fmt = para.paragraph_format
    fmt.space_before = Emu(_gap_to_emu(prev_block, block))
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(max(_max_span_size(block) * 1.15, 6.0))
    for span in block["spans"]:
        run = para.add_run(span["text"])
        run.bold = span.get("bold", False)
        run.italic = span.get("italic", False)
        run.font.size = Pt(span.get("size", 11))
        run.font.color.rgb = _int_color_to_rgb(span.get("color", 0))

def _add_image_block(container, page, block, prev_block=None):
    para = container.add_paragraph()
    para.paragraph_format.space_before = Emu(_gap_to_emu(prev_block, block))
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    width_in, height_in = _image_inches_for_page(page, block)
    run.add_picture(io.BytesIO(block["data"]), width=Inches(width_in), height=Inches(height_in))

def _add_side_by_side_table(docx, page, img_block, text_blocks, prev_block=None):
    table = docx.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)
    row = table.rows[0]
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    img_first = img_block["bbox"][0] <= min(tb["bbox"][0] for tb in text_blocks)
    img_cell = row.cells[0] if img_first else row.cells[1]
    text_cell = row.cells[1] if img_first else row.cells[0]

    gap_pt = max(12.0, min(abs(tb["bbox"][0] - img_block["bbox"][2]) for tb in text_blocks))
    content_width_pt = max(72.0, page.rect.width - docx.sections[0].left_margin.pt - docx.sections[0].right_margin.pt)
    img_width_pt = min(img_block["width_pt"], content_width_pt * 0.65)
    text_width_pt = max(72.0, content_width_pt - img_width_pt - gap_pt)
    _set_cell_width(img_cell, img_width_pt)
    _set_cell_width(text_cell, text_width_pt)

    img_para = img_cell.paragraphs[0]
    img_para.paragraph_format.space_before = Emu(_gap_to_emu(prev_block, img_block))
    img_para.paragraph_format.space_after = Pt(0)
    width_in, height_in = _image_inches_for_page(page, {**img_block, "width_pt": img_width_pt})
    img_para.add_run().add_picture(io.BytesIO(img_block["data"]), width=Inches(width_in), height=Inches(height_in))

    text_blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
    prev_text = None
    first_para = text_cell.paragraphs[0]
    first_para._element.getparent().remove(first_para._element)
    for tb in text_blocks:
        _add_text_block(text_cell, tb, prev_text)
        prev_text = tb

def _add_page_snapshot(docx, page, prev_block=None):
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), colorspace=fitz.csRGB)
    block = {
        "kind": "image",
        "bbox": (0.0, 0.0, float(page.rect.width), float(page.rect.height)),
        "data": pix.tobytes("png"),
        "width_pt": float(page.rect.width),
        "height_pt": float(page.rect.height),
    }
    _add_image_block(docx, page, block, prev_block)

def _image_inches_for_page(page, block):
    max_width_pt = max(36.0, page.rect.width)
    width_pt = min(float(block["width_pt"]), max_width_pt)
    ratio = width_pt / max(float(block["width_pt"]), 1.0)
    height_pt = float(block["height_pt"]) * ratio
    return width_pt / PT_PER_INCH, height_pt / PT_PER_INCH

def _max_span_size(block):
    return max((float(s.get("size", 11)) for s in block.get("spans", [])), default=11.0)

def _int_color_to_rgb(color_int):
    return RGBColor((int(color_int) >> 16) & 0xFF, (int(color_int) >> 8) & 0xFF, int(color_int) & 0xFF)

def _union_blocks(blocks):
    return {
        "kind": "group",
        "bbox": (
            min(b["bbox"][0] for b in blocks),
            min(b["bbox"][1] for b in blocks),
            max(b["bbox"][2] for b in blocks),
            max(b["bbox"][3] for b in blocks),
        )
    }

def _remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{name}"))
        if border is None:
            border = OxmlElement(f"w:{name}")
            borders.append(border)
        border.set(qn("w:val"), "nil")

def _set_cell_width(cell, width_pt):
    width_twips = str(int(width_pt * 20))
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), width_twips)
    tc_w.set(qn("w:type"), "dxa")

def _fix_docx_margins(docx_path, pdf_bytes):
    """
    Match DOCX page layout to the source PDF:
    - Page margins (L/T/R/B) from content bounding box
    - header_distance matched to first text line in PDF
    - paragraph space_before/space_after on first N paragraphs matched
      to the exact inter-line gaps in the PDF header zone
    """
    try:
        doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
        if doc_pdf.page_count == 0:
            doc_pdf.close(); return
        page = doc_pdf[0]
        pw, ph = page.rect.width, page.rect.height

        # ── All text blocks ─────────────────────────────────────────────
        blocks = page.get_text("dict")["blocks"]
        text_blocks = sorted(
            [b for b in blocks if b.get("type") == 0],
            key=lambda b: b["bbox"][1]
        )
        if not text_blocks:
            doc_pdf.close(); return

        # ── Page margins from content bounds ─────────────────────────────
        all_x = [b["bbox"][0] for b in text_blocks] + [b["bbox"][2] for b in text_blocks]
        all_y = [b["bbox"][1] for b in text_blocks] + [b["bbox"][3] for b in text_blocks]
        L = max(0.0, min(all_x))
        T = max(0.0, min(all_y))
        R = max(0.0, pw - max(all_x))
        B = max(0.0, ph - max(all_y))

        # ── Header zone lines (top 15 %) ──────────────────────────────────
        hdr_limit = ph * 0.15
        hdr_lines = []
        for blk in text_blocks:
            if blk["bbox"][1] >= hdr_limit:
                break
            for line in blk.get("lines", []):
                bb = line.get("bbox", [0, 0, 0, 0])
                spans = line.get("spans", [])
                fs = max((s.get("size", 10) for s in spans), default=10)
                hdr_lines.append({"y0": bb[1], "y1": bb[3], "fs": fs})
        hdr_lines.sort(key=lambda x: x["y0"])

        doc_pdf.close()

        # ── Apply to DOCX ────────────────────────────────────────────────
        docx = DocxDocument(docx_path)

        # 1. Page margins + header distance
        for section in docx.sections:
            section.left_margin   = Pt(L)
            section.top_margin    = Pt(T)
            section.right_margin  = Pt(R)
            section.bottom_margin = Pt(B)
            if hdr_lines:
                # header_distance = distance from page top to first header text
                section.header_distance = Pt(max(5.0, hdr_lines[0]["y0"]))

        # 2. Fix paragraph spacing on first paragraphs to match PDF gaps
        paras = [p for p in docx.paragraphs]
        n_hdr = len(hdr_lines)
        for i, para in enumerate(paras[:max(n_hdr + 2, 4)]):
            fmt = para.paragraph_format
            fmt.space_before = Pt(0)           # remove all leading space
            if i < n_hdr:
                ld = hdr_lines[i]
                if i + 1 < n_hdr:
                    gap = hdr_lines[i + 1]["y0"] - ld["y1"]
                    fmt.space_after = Pt(max(0.0, gap))
                else:
                    fmt.space_after = Pt(0)
                # Pin line spacing to exact font size (single line, no expansion)
                fmt.line_spacing = Pt(ld["fs"] * 1.2)

        docx.save(docx_path)
        print(f"[DOCX] Margins L={L:.0f} T={T:.0f} R={R:.0f} B={B:.0f} pt  "
              f"hdr_lines={n_hdr}  hdr_dist={hdr_lines[0]['y0']:.0f}pt" if hdr_lines else
              f"[DOCX] Margins L={L:.0f} T={T:.0f} R={R:.0f} B={B:.0f} pt")
    except Exception as e:
        print(f"[DOCX] Margin fix failed: {e}")
        traceback.print_exc()

# ─────────────────────────────────────────────────────────────────────────────
# Header / footer detection & real Word header/footer application
# ─────────────────────────────────────────────────────────────────────────────

def _extract_header_footer_info(pdf_bytes):
    """
    Detect a consistent header/footer across pages, tolerating a changing
    page number inside the footer (e.g. "Page 1", "Page 2", ...).
    Returns (header_info, footer_info) where each is a dict:
      {"text": str, "font_size": float, "bold": bool, "template": str,
       "has_page_number": bool}
    or None if no consistent header/footer was found.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    n   = doc.page_count
    if n == 0:
        doc.close(); return None, None

    page_headers = []
    page_footers = []

    for page in doc:
        h  = page.rect.height
        w  = page.rect.width

        # Header zone: top 10% of page
        hdr_rect = fitz.Rect(0, 0, w, h * 0.10)
        # Footer zone: bottom 10% of page
        ftr_rect = fitz.Rect(0, h * 0.90, w, h)

        hdr_blocks = page.get_text("dict", clip=hdr_rect)["blocks"]
        ftr_blocks = page.get_text("dict", clip=ftr_rect)["blocks"]

        page_headers.append(_blocks_to_info(hdr_blocks))
        page_footers.append(_blocks_to_info(ftr_blocks))

    doc.close()

    hdr_info = _pick_dominant(page_headers, n, allow_page_number=False)
    ftr_info = _pick_dominant(page_footers, n, allow_page_number=True)
    return hdr_info, ftr_info

def _blocks_to_info(blocks):
    """Summarise text blocks into {text, lines, font_size, bold}.

    `lines` preserves the individual PDF line breaks (as opposed to `text`,
    which joins everything with spaces) because pdf2docx / our own repair
    step reconstruct the header as separate Word paragraphs — one per
    original PDF line, or one per table cell for multi-column headers — not
    as a single joined paragraph. Keeping the per-line breakdown lets the
    dedup step match and remove each of those individual paragraphs.
    """
    lines  = []
    sizes  = []
    bolds  = []
    for blk in blocks:
        if blk.get("type") != 0: continue
        for line in blk.get("lines", []):
            txt = "".join(s.get("text","") for s in line.get("spans",[]))
            if txt.strip():
                lines.append(txt.strip())
                for span in line.get("spans", []):
                    sizes.append(span.get("size", 11))
                    bolds.append(bool(span.get("flags", 0) & 0b10000))
    if not lines:
        return None
    text      = " ".join(lines)
    font_size = round(sum(sizes)/len(sizes), 1) if sizes else 11.0
    bold      = sum(bolds) > len(bolds) / 2
    return {"text": text, "lines": lines, "font_size": font_size, "bold": bold}

def _norm_dynamic(text):
    """Normalize digit runs so 'Page 1' and 'Page 2' compare as the same
    template — lets us detect a stable footer even when the page number
    inside it changes on every page."""
    return re.sub(r'\d+', '#', text or "")

def _detect_incrementing(infos):
    """True if the first digit run increases by ~1 across consecutive pages
    that have this info — a strong signal it's a live page number rather
    than static text that happens to contain a digit."""
    nums = []
    for info in infos:
        if not info:
            continue
        m = re.search(r'\d+', info["text"])
        if m:
            nums.append(int(m.group()))
    if len(nums) < 2:
        return False
    diffs = [b - a for a, b in zip(nums, nums[1:])]
    return sum(1 for d in diffs if d == 1) >= max(1, len(diffs) - 1)

def _pick_dominant(infos, n_pages, allow_page_number=False):
    """
    Return the most common info, matching on a digit-normalized template so
    that a changing page number ('Page 1' vs 'Page 2') doesn't prevent a
    footer from being detected as consistent across pages.
    """
    from collections import Counter
    norm_texts = [_norm_dynamic(i["text"]) if i else "" for i in infos]
    cnt = Counter(norm_texts)
    top_norm, top_count = cnt.most_common(1)[0]
    if not top_norm or top_count < max(2, n_pages * 0.4):
        return None

    for info in infos:
        if info and _norm_dynamic(info["text"]) == top_norm:
            result = dict(info)
            result["template"] = top_norm
            # Per-line templates let the dedup step match and strip each
            # individual fragment/table-cell paragraph pdf2docx produced,
            # not just a single paragraph containing the whole joined text
            # (which usually doesn't exist as one paragraph at all).
            result["line_templates"] = [_norm_dynamic(l) for l in info.get("lines", [])]
            result["has_page_number"] = (
                allow_page_number and _detect_incrementing(infos)
            )
            return result
    return None

# ── Word field helpers (live PAGE / NUMPAGES instead of frozen numbers) ─────

def _add_field(paragraph, field_code, placeholder="1"):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_code} '
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t'); fld_text.text = placeholder
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r = run._r
    for el in (fld_begin, instr, fld_sep, fld_text, fld_end):
        r.append(el)

_PAGE_OF_TOTAL_RE = re.compile(r'(\d+)(\s*(?:of|/)\s*)(\d+)', re.I)

def _write_hf_text_with_fields(paragraph, text, use_fields):
    """
    Write text into a header/footer paragraph. If use_fields is True and the
    text contains a page-number pattern, swap it for a live PAGE / NUMPAGES
    field so it updates correctly per page in Word instead of staying frozen
    at whatever number the source PDF happened to show.
    """
    if not use_fields:
        paragraph.add_run(text)
        return

    m = _PAGE_OF_TOTAL_RE.search(text)
    if m:
        before, sep, after = text[:m.start()], m.group(2), text[m.end():]
        if before: paragraph.add_run(before)
        _add_field(paragraph, "PAGE", m.group(1))
        paragraph.add_run(sep)
        _add_field(paragraph, "NUMPAGES", m.group(3))
        if after: paragraph.add_run(after)
        return

    m = re.search(r'\d+', text)
    if m:
        before, after = text[:m.start()], text[m.end():]
        if before: paragraph.add_run(before)
        _add_field(paragraph, "PAGE", m.group())
        if after: paragraph.add_run(after)
        return

    paragraph.add_run(text)

def _set_hf_paragraph(hf_section, info):
    """Set text, font size and bold on the first paragraph of a
    header/footer, using live page-number fields when appropriate."""
    for para in hf_section.paragraphs:
        for run in para.runs:
            run.text = ""

    if hf_section.paragraphs:
        para = hf_section.paragraphs[0]
    else:
        para = hf_section.add_paragraph()

    para.clear()
    _write_hf_text_with_fields(para, info["text"], info.get("has_page_number", False))
    if para.runs:
        para.runs[0].bold      = info.get("bold", False)
        para.runs[0].font.size = Pt(info.get("font_size", 11))

def _apply_headers_footers(docx_path, hdr_info, ftr_info):
    """
    Write hdr_info/ftr_info into real Word header/footer sections, and strip
    the duplicated per-page copies pdf2docx already placed in the body.

    Matching happens against BOTH the whole-block template and each
    individual line's template, and scans EVERY paragraph in the body
    including ones nested inside tables — pdf2docx frequently rebuilds a
    multi-column PDF header (e.g. "Subject Code | Subject Name" and
    "Enrollment No | Name" rows) as a Word table, and our header-wrap repair
    reconstructs one paragraph per original PDF line rather than a single
    paragraph for the whole header. Only checking `document.paragraphs`
    (top-level only) or only the fully-joined text would miss these and
    leave the duplicate sitting in the body, which is what was happening.
    """
    doc = DocxDocument(docx_path)

    templates_to_remove = set()
    for info in (hdr_info, ftr_info):
        if not info:
            continue
        templates_to_remove.add(info["template"])
        templates_to_remove.update(info.get("line_templates", []))
    templates_to_remove.discard("")

    paras_to_delete = [
        para for para in _iter_body_paragraphs(doc)
        if para.text.strip() and _norm_dynamic(para.text.strip()) in templates_to_remove
    ]
    for para in paras_to_delete:
        p = para._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)
    print(f"[DOCX] Removed {len(paras_to_delete)} duplicate header/footer paragraph(s) from body")

    # A table that pdf2docx used purely to lay out the header/footer will
    # now be fully empty — remove it so no stray blank box is left behind.
    _remove_empty_tables(doc)

    for section in doc.sections:
        section.different_first_page_header_footer = False
        if hdr_info:
            _set_hf_paragraph(section.header, hdr_info)
        if ftr_info:
            _set_hf_paragraph(section.footer, ftr_info)

    doc.save(docx_path)
    print(f"[DOCX] Applied header={hdr_info} footer={ftr_info}")

@app.route("/download/<token>")
def download(token):
    item = _store_pop(token)
    if not item:
        return jsonify({"error": "File not found or already downloaded"}), 404
    return send_file(
        io.BytesIO(item["data"]),
        mimetype=item["mime"],
        as_attachment=True,
        download_name=item["filename"]
    )

# ─────────────────────────────────────────────────────────────────────────────
# PDF Edit — change header dates & footer page numbers
# ─────────────────────────────────────────────────────────────────────────────

_DATE_PATTERNS = [
    re.compile(r'\d{1,2}\s*[/\-\.]\s*\d{1,2}\s*[/\-\.]\s*\d{2,4}'),  # 20-09-2025 / 20 / 09 / 25
    re.compile(r'\d{4}\s*[/\-\.]\s*\d{1,2}\s*[/\-\.]\s*\d{1,2}'),     # 2025-09-20
    re.compile(r'\d{1,2}\s+[\w]{3,9},?\s+\d{2,4}', re.I),             # 20 September 2025
    re.compile(r'\d{1,2}\s*[/\-\.]\s*[A-Za-z]{3,9}\s*[/\-\.]\s*\d{2,4}', re.I), # 20-Jan-2025
    re.compile(r'[A-Za-z]{3,9}\s+\d{1,2},?\s+\d{2,4}', re.I),         # Jan 20, 2025
]
_TRAILING_PAGE_DIGITS = re.compile(r'(\d{1,4})\s*$')
_ALL_PAGE_ALIASES = {
    "all", "allpage", "allpages", "all page", "all pages",
    "every", "everypage", "everypages", "every page", "every pages", "*",
}

def _parse_page_range(s, max_pages):
    """'1-5,7' → sorted list of 0-based page indices. Blank/all = all pages."""
    s = str(s or "").strip().lower()
    compact = re.sub(r'[\s_\-]+', '', s)
    if not s or s in _ALL_PAGE_ALIASES or compact in _ALL_PAGE_ALIASES:
        return list(range(max_pages))
    pages = set()
    for part in re.split(r'[,;]', s):
        part = part.strip()
        if not part:
            continue
        part_compact = re.sub(r'[\s_\-]+', '', part)
        if part in _ALL_PAGE_ALIASES or part_compact in _ALL_PAGE_ALIASES:
            pages.update(range(max_pages))
            continue
        m = re.match(r'^(\d+)\s*-\s*(\d+)$', part)
        if m:
            a, b = int(m.group(1)), int(m.group(2))
            if a > b:
                a, b = b, a
            pages.update(range(max(0, a - 1), min(max_pages, b)))
        elif part.isdigit():
            p = int(part) - 1
            if 0 <= p < max_pages:
                pages.add(p)
    return sorted(pages)

def _redact_and_insert(page, doc, rect, new_text, span):
    """Redact old text region and insert new_text with same font/size/color."""
    # Sample background
    fill = (1.0, 1.0, 1.0)
    try:
        pm = page.get_pixmap(clip=rect.irect, matrix=fitz.Matrix(1, 1), colorspace=fitz.csRGB)
        if pm.samples:
            fill = (pm.samples[0]/255.0, pm.samples[1]/255.0, pm.samples[2]/255.0)
    except: pass

    page.add_redact_annot(rect, fill=fill)
    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    chars   = span.get("chars", [])
    if not chars: return False
    fs      = float(span.get("size", 10))
    color   = _unpack_color(int(span.get("color", 0)))
    fn      = span.get("font", "helv")
    flags   = int(span.get("flags", 0))
    origin  = fitz.Point(chars[0]["origin"][0], chars[0]["origin"][1])

    font_obj, is_real = _load_font(doc, page, fn, flags)
    if is_real and font_obj and _font_covers_text(font_obj, new_text):
        try:
            tw = fitz.TextWriter(page.rect)
            tw.append(origin, new_text, font=font_obj, fontsize=fs)
            tw.write_text(page, color=color)
            return True
        except: pass
    page.insert_text(origin, new_text,
                     fontname=_pick_builtin(fn, flags),
                     fontsize=fs, color=color)
    return True

def _line_chars_from_rawdict(raw):
    """Return page text lines as character records with span metadata attached."""
    lines = []
    for blk in raw.get("blocks", []):
        if blk.get("type") != 0:
            continue
        for line in blk.get("lines", []):
            chars = []
            for span in line.get("spans", []):
                meta = {
                    "font": span.get("font", "helv"),
                    "size": float(span.get("size", 10)),
                    "color": int(span.get("color", 0)),
                    "flags": int(span.get("flags", 0)),
                }
                for ch in span.get("chars", []):
                    if "bbox" not in ch:
                        continue
                    chars.append({
                        "c": ch.get("c", ""),
                        "bbox": ch["bbox"],
                        "origin": ch.get("origin", (ch["bbox"][0], ch["bbox"][3])),
                        "span": meta,
                    })
            if chars:
                lines.append(chars)
    lines.sort(key=lambda cs: (min(c["bbox"][1] for c in cs), min(c["bbox"][0] for c in cs)))
    return lines

def _rect_from_chars(chars, pad_x=1.5, pad_y=1.0):
    return fitz.Rect(
        min(c["bbox"][0] for c in chars) - pad_x,
        min(c["bbox"][1] for c in chars) - pad_y,
        max(c["bbox"][2] for c in chars) + pad_x,
        max(c["bbox"][3] for c in chars) + pad_y,
    )

def _sample_rect_fill(page, rect):
    fill = (1.0, 1.0, 1.0)
    try:
        pm = page.get_pixmap(clip=rect.irect, matrix=fitz.Matrix(1, 1), colorspace=fitz.csRGB)
        if pm.samples and len(pm.samples) >= 3:
            fill = (pm.samples[0] / 255, pm.samples[1] / 255, pm.samples[2] / 255)
    except Exception:
        pass
    return fill

def _insert_like_chars(page, doc, chars, new_text, color_override=None):
    if not chars:
        return False
    first = chars[0]
    span = first["span"]
    fs = float(span.get("size", 10))
    fn = span.get("font", "helv")
    flags = int(span.get("flags", 0))
    color = color_override if color_override is not None else _unpack_color(int(span.get("color", 0)))
    origin = fitz.Point(first["origin"][0], first["origin"][1])
    rect = _rect_from_chars(chars)

    font_obj, is_real = _load_font(doc, page, fn, flags)
    use_embedded = is_real and font_obj and _font_covers_text(font_obj, new_text)
    for size in [fs, fs * 0.96, fs * 0.92, fs * 0.88, fs * 0.84, fs * 0.80]:
        if use_embedded:
            try:
                tw = fitz.TextWriter(page.rect)
                tw.append(origin, new_text, font=font_obj, fontsize=size)
                tw.write_text(page, color=color)
                return True
            except Exception:
                pass
        try:
            page.insert_text(origin, new_text,
                             fontname=_pick_builtin(fn, flags),
                             fontsize=size, color=color)
            return True
        except Exception:
            pass

    try:
        page.insert_textbox(rect, new_text,
                            fontname=_pick_builtin(fn, flags),
                            fontsize=max(5.0, fs * 0.75), color=color)
        return True
    except Exception:
        return False

def _replace_header_dates_on_page(page, doc, new_date):
    h = page.rect.height
    w = page.rect.width
    # Header first; if not found, fall back to the full page because some PDFs
    # store header text outside the visible top band or with unusual coordinates.
    count = _replace_dates_in_rect(page, doc, fitz.Rect(0, 0, w, h * 0.28), new_date)
    if count:
        return count
    return _replace_dates_in_rect(page, doc, page.rect, new_date)

def _replace_dates_in_rect(page, doc, rect, new_date):
    raw = page.get_text("rawdict", clip=rect, flags=fitz.TEXT_PRESERVE_WHITESPACE)
    replacements = []
    for chars in _line_chars_from_rawdict(raw):
        text = "".join(c["c"] for c in chars)
        for pattern in _DATE_PATTERNS:
            for match in pattern.finditer(text):
                match_chars = chars[match.start():match.end()]
                if match_chars:
                    replacements.append(match_chars)

    if not replacements:
        return 0

    fills = []
    for chars in replacements:
        rect = _rect_from_chars(chars)
        fills.append((rect, _sample_rect_fill(page, rect), chars))

    for rect, fill, _ in fills:
        page.add_redact_annot(rect, fill=fill)
    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    count = 0
    for _, _, chars in fills:
        if _insert_like_chars(page, doc, chars, new_date):
            count += 1
    return count

def _edit_pdf(doc, date_rules, start_page_num):
    """
    Apply date replacements in header zones and page-number
    renumbering in footer zones.
    Returns total number of changes.
    """
    n = doc.page_count
    changes = 0

    # Build page → new_date map
    page_date = {}
    for rule in date_rules:
        pages_str = str(rule.get("pages", "")).strip()
        date_str  = str(rule.get("date",  "")).strip()
        if not date_str: continue
        if pages_str:
            for p in _parse_page_range(pages_str, n):
                if p not in page_date:
                    page_date[p] = date_str
        else:
            # Blank range = all pages
            for p in range(n):
                if p not in page_date:
                    page_date[p] = date_str

    for i, page in enumerate(doc):
        h = page.rect.height
        w = page.rect.width
        hdr_zone = fitz.Rect(0, 0, w, h * 0.15)

        # ── Header: replace date ─────────────────────────────────────────
        if i in page_date:
            new_date = page_date[i]
            try:
                changes += _replace_header_dates_on_page(page, doc, new_date)
            except Exception:
                traceback.print_exc()

        # ── Footer: renumber page ────────────────────────────────────────
        if start_page_num is not None:
            new_pg = str(start_page_num + i)
            if _replace_page_number(page, doc, new_pg):
                changes += 1

    return changes

def _replace_page_number(page, doc, new_pg):
    """
    Replace the trailing 1-4 digits on the bottom-most footer text line.
    This handles bare numbers and labels such as 'Page No: 0012'.
    """
    try:
        h = page.rect.height
        w = page.rect.width
        footer_zone = fitz.Rect(0, h * 0.55, w, h)
        raw = page.get_text("rawdict", clip=footer_zone,
                            flags=fitz.TEXT_PRESERVE_WHITESPACE)
    except Exception as e:
        print(f"[EDIT] Page {page.number+1}: footer text failed: {e}")
        return False

    lines = sorted(_line_chars_from_rawdict(raw),
                   key=lambda cs: (max(c["bbox"][3] for c in cs), max(c["bbox"][2] for c in cs)),
                   reverse=True)
    candidate = None
    for chars in lines:
        text = "".join(c["c"] for c in chars)
        m = _TRAILING_PAGE_DIGITS.search(text)
        if not m:
            continue
        digit_chars = chars[m.start(1):m.end(1)]
        if digit_chars:
            candidate = (text, digit_chars)
            break

    if not candidate:
        # Fallback: choose the right-most 1-4 digit group on one of the lowest
        # text lines. This catches PDFs where labels and numbers are extracted
        # in odd order, while still avoiding body text.
        for chars in lines[:8]:
            text = "".join(c["c"] for c in chars)
            matches = list(re.finditer(r'\d{1,4}', text))
            if not matches:
                continue
            best = max(matches, key=lambda m: max(c["bbox"][2] for c in chars[m.start():m.end()]))
            digit_chars = chars[best.start():best.end()]
            if digit_chars:
                candidate = (text, digit_chars)
                break

    if not candidate:
        print(f"[EDIT] Page {page.number+1}: no trailing footer digits found")
        return False

    old_line, digit_chars = candidate
    old_num = "".join(c["c"] for c in digit_chars).strip()
    exp = _rect_from_chars(digit_chars, pad_x=2.0, pad_y=1.5)
    fill = _sample_rect_fill(page, exp)
    page.add_redact_annot(exp, fill=fill)
    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    ok = _insert_page_digits_like_chars(page, digit_chars, str(int(new_pg)))
    if ok:
        print(f"[EDIT] Page {page.number+1}: trailing footer '{old_num}' → '{new_pg}' in '{old_line.strip()}'")
    return ok

def _insert_page_digits_like_chars(page, chars, new_text):
    """Write page digits without leading zeroes, right-aligned in old digit box."""
    if not chars:
        return False
    span = chars[-1]["span"]
    fs = float(span.get("size", 10))
    fn = span.get("font", "helv")
    flags = int(span.get("flags", 0))
    color = _unpack_color(int(span.get("color", 0)))
    rect = _rect_from_chars(chars, pad_x=1.5, pad_y=1.5)
    fontname = _pick_builtin(fn, flags)

    for size in [fs, fs * 0.96, fs * 0.92, fs * 0.88, fs * 0.84, fs * 0.80]:
        try:
            text_width = fitz.get_text_length(new_text, fontname=fontname, fontsize=max(5.0, size))
            x = max(rect.x0, rect.x1 - text_width)
            origin = fitz.Point(x, chars[-1]["origin"][1])
            page.insert_text(origin, new_text, fontname=fontname, fontsize=max(5.0, size), color=color)
            return True
        except Exception:
            pass

    try:
        origin = fitz.Point(chars[-1]["bbox"][2] - max(6.0, fs * len(new_text) * 0.55),
                            chars[-1]["origin"][1])
        page.insert_text(origin, new_text, fontname=fontname, fontsize=fs, color=color)
        return True
    except Exception:
        return False

@app.route("/pdf-edit", methods=["POST"])
def pdf_edit():
    """Edit PDF: change header dates by page range, renumber footer pages."""
    import json as _json
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f = request.files["file"]
    if not f.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Must be a PDF"}), 400

    # Date rules
    try:
        date_rules = _json.loads(request.form.get("date_rules", "[]"))
    except:
        date_rules = []

    # Start page number
    spn_raw = request.form.get("start_page_num", "").strip()
    start_page_num = int(spn_raw) if spn_raw.isdigit() else None

    output_name = request.form.get("output_name", "").strip()

    if not date_rules and start_page_num is None:
        return jsonify({"error": "Provide at least one date rule or a starting page number"}), 400

    pdf_bytes = f.read()
    doc       = fitz.open(stream=pdf_bytes, filetype="pdf")
    changes   = _edit_pdf(doc, date_rules, start_page_num)

    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True)
    doc.close()
    buf.seek(0)

    base     = os.path.splitext(f.filename)[0]
    out_stem = re.sub(r'[^\w\s\-]', '', output_name).strip() if output_name else f"{base}_edited"
    token    = _store_put(buf.getvalue(), "application/pdf", f"{out_stem}.pdf")
    return jsonify({"token": token, "changes": changes,
                    "filename": f"{out_stem}.pdf"})

@app.route("/pdf-debug", methods=["POST"])
def pdf_debug():
    """Return all text spans per page with positions — for debugging footer detection."""
    if "file" not in request.files:
        return jsonify({"error": "No file"}), 400
    f = request.files["file"]
    pdf_bytes = f.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    result = []
    for i, page in enumerate(doc):
        h = page.rect.height
        w = page.rect.width
        footer_zone = fitz.Rect(0, h * 0.85, w, h)
        page_info = {"page": i+1, "height": h, "width": w,
                     "footer_y_start": round(h * 0.85, 1),
                     "footer_spans": [], "all_bottom_spans": []}
        # Footer zone spans
        try:
            blocks = page.get_text("dict", clip=footer_zone).get("blocks", [])
            for blk in blocks:
                if blk.get("type") != 0: continue
                for line in blk.get("lines", []):
                    for span in line.get("spans", []):
                        bbox = span.get("bbox", [])
                        page_info["footer_spans"].append({
                            "text": repr(span.get("text","")),
                            "y0": round(bbox[1],1) if len(bbox)>1 else "?",
                            "y1": round(bbox[3],1) if len(bbox)>3 else "?",
                        })
        except Exception as e:
            page_info["footer_error"] = str(e)
        # Bottom 10 spans on whole page
        try:
            all_spans = []
            blocks2 = page.get_text("dict").get("blocks", [])
            for blk in blocks2:
                if blk.get("type") != 0: continue
                for line in blk.get("lines", []):
                    for span in line.get("spans", []):
                        bbox = span.get("bbox", [])
                        if len(bbox) >= 4:
                            all_spans.append((bbox[3], repr(span.get("text","")), round(bbox[1],1), round(bbox[3],1)))
            all_spans.sort(key=lambda x: -x[0])
            page_info["all_bottom_spans"] = [
                {"text": t, "y0": y0, "y1": y1} for _, t, y0, y1 in all_spans[:10]
            ]
        except: pass
        result.append(page_info)
        if i >= 2: break   # Only check first 3 pages
    doc.close()
    return jsonify(result)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/status")
def status():
    return jsonify({
        "pdf2docx": PDF2DOCX_AVAILABLE,
        "pymupdf":  fitz.__version__,
    })

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5055, debug=False)
